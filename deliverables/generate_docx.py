"""Direct markdown -> docx converter for the RFQ2BOQ internship report.
Handles: h1/h2/h3, paragraphs with **bold**/*italic*, bullet/numbered lists,
pipe-tables, horizontal rules, page-break markers, and the two custom
signature-block HTML tables (parsed specially, not via generic HTML).
"""
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1]
OUT = sys.argv[2]
SIG_IMAGE = sys.argv[3]
CERT_IMAGE = sys.argv[4] if len(sys.argv) > 4 else None
BUDDY_SIG_IMAGE = sys.argv[5] if len(sys.argv) > 5 else None

with open(SRC, "r", encoding="utf-8") as f:
    text = f.read()

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(11)

HEADING_COLOR = RGBColor(0x0F, 0x1E, 0x30)


def add_page_break():
    doc.add_page_break()


def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADING_COLOR
        run.font.name = "Georgia"
    return h


def add_runs_with_bold(paragraph, raw):
    """Split on **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*[^*]+\*\*)", raw)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # handle simple *italic*
            subparts = re.split(r"(\*[^*]+\*)", part)
            for sp in subparts:
                if not sp:
                    continue
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    r = paragraph.add_run(sp[1:-1])
                    r.italic = True
                else:
                    paragraph.add_run(sp)


def add_paragraph(raw, style_name=None):
    p = doc.add_paragraph(style=style_name)
    add_runs_with_bold(p, raw)
    return p


def set_cell_text(cell, raw, bold_label=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if "@@SIGIMG@@" in raw and SIG_IMAGE:
        before, _, after = raw.partition("@@SIGIMG@@")
        if before.strip():
            add_runs_with_bold(p, before)
        run = p.add_run()
        run.add_picture(SIG_IMAGE, width=Inches(1.6))
        if after.strip():
            add_runs_with_bold(p, after)
        return
    if "@@BUDDYSIGIMG@@" in raw and BUDDY_SIG_IMAGE:
        before, _, after = raw.partition("@@BUDDYSIGIMG@@")
        if before.strip():
            add_runs_with_bold(p, before)
        run = p.add_run()
        run.add_picture(BUDDY_SIG_IMAGE, width=Inches(1.8))
        if after.strip():
            add_runs_with_bold(p, after)
        return
    add_runs_with_bold(p, raw)


def add_signature_table(rows):
    """rows: list of (col1, col2_or_None)"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # remove borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)
    for i, (c1, c2) in enumerate(rows):
        set_cell_text(table.cell(i, 0), c1)
        if c2 is not None:
            set_cell_text(table.cell(i, 1), c2)
        else:
            table.cell(i, 0).merge(table.cell(i, 1))
    doc.add_paragraph()


# --- Extract and remove the two custom HTML signature blocks, replacing
# with placeholder tokens we handle natively ---

def extract_table_rows(html_table):
    rows = []
    for row_match in re.finditer(r"<tr>(.*?)</tr>", html_table, re.DOTALL):
        row_html = row_match.group(1)
        cells = re.findall(r"<td[^>]*>(.*?)</td>", row_html, re.DOTALL)
        cells = [re.sub(r"<[^>]+>", "", c).strip() for c in cells]
        cells = [c.replace("&nbsp;", " ") for c in cells]
        if len(cells) == 1:
            rows.append((cells[0], None))
        elif len(cells) >= 2:
            rows.append((cells[0], cells[1]))
    return rows


# Replace [SIGNATURE_IMAGE] marker before table extraction
text = text.replace("[SIGNATURE_IMAGE]", "@@SIGIMG@@")
text = text.replace("[BUDDY_SIGNATURE_IMAGE]", "@@BUDDYSIGIMG@@")
text = text.replace("[CERTIFICATE_IMAGE]", "<<<CERTIMAGE>>>")

html_tables = re.findall(r"<table[^>]*>.*?</table>", text, re.DOTALL)
table_data = [extract_table_rows(t) for t in html_tables]
for t in html_tables:
    text = text.replace(t, "\n<<<SIGTABLE>>>\n")

text = re.sub(r'<div style="height:\d+px;"></div>', "", text)
text = re.sub(r'<p style="text-align:center;">(.*?)</p>', r"\n<<<CENTER>>>\1\n", text, flags=re.DOTALL)
text = re.sub(r'<div style="page-break-before: always;"></div>', "\n<<<PAGEBREAK>>>\n", text)

lines = text.split("\n")

i = 0
sig_table_idx = 0
in_table = False
table_buffer = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if stripped == "<<<PAGEBREAK>>>":
        add_page_break()
        i += 1
        continue

    if stripped == "<<<SIGTABLE>>>":
        rows = table_data[sig_table_idx]
        sig_table_idx += 1
        add_signature_table(rows)
        i += 1
        continue

    if stripped == "<<<CERTIMAGE>>>":
        if CERT_IMAGE:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(CERT_IMAGE, width=Inches(4.4))
        i += 1
        continue

    if stripped.startswith("<<<CENTER>>>"):
        content = stripped.replace("<<<CENTER>>>", "")
        content = content.replace("<b>", "**").replace("</b>", "**")
        content = content.replace("<i>", "*").replace("</i>", "*")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        segments = content.split("<br>")
        for idx, seg in enumerate(segments):
            if idx > 0:
                p.add_run().add_break()
            add_runs_with_bold(p, seg)
        i += 1
        continue

    if stripped == "---":
        i += 1
        continue

    if stripped == "":
        i += 1
        continue

    if stripped.startswith("# "):
        add_heading(stripped[2:].strip(), 1)
        i += 1
        continue
    if stripped.startswith("## "):
        add_heading(stripped[3:].strip(), 2)
        i += 1
        continue
    if stripped.startswith("### "):
        add_heading(stripped[4:].strip(), 3)
        i += 1
        continue

    # pipe tables
    if stripped.startswith("|"):
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            row_line = lines[i].strip()
            if re.match(r"^\|[\s:\-|]+\|$", row_line):
                i += 1
                continue
            cells = [c.strip() for c in row_line.strip("|").split("|")]
            table_rows.append(cells)
            i += 1
        if table_rows:
            ncols = len(table_rows[0])
            t = doc.add_table(rows=len(table_rows), cols=ncols)
            t.style = "Light Grid Accent 1"
            for r_idx, row in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < ncols:
                        set_cell_text(t.cell(r_idx, c_idx), cell_text)
            doc.add_paragraph()
        continue

    # bullet list
    if stripped.startswith("- ") or stripped.startswith("* "):
        p = add_paragraph(stripped[2:], style_name="List Bullet")
        i += 1
        continue

    # numbered list
    m = re.match(r"^(\d+)\.\s+(.*)", stripped)
    if m:
        p = add_paragraph(m.group(2), style_name="List Number")
        i += 1
        continue

    # code fence - skip fenced blocks, render as monospace paragraph
    if stripped.startswith("```"):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i])
            i += 1
        i += 1  # skip closing fence
        p = doc.add_paragraph("\n".join(code_lines))
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        continue

    # regular paragraph (collect continuation lines until blank)
    para_lines = [stripped]
    i += 1
    while i < len(lines) and lines[i].strip() != "" and not lines[i].strip().startswith(("#", "|", "-", "*", "<<<", "```")) and not re.match(r"^\d+\.", lines[i].strip()):
        para_lines.append(lines[i].strip())
        i += 1
    add_paragraph(" ".join(para_lines))

doc.save(OUT)
print(f"Saved {OUT}")
